#!/usr/bin/env python3
"""
convert.py — PDF conversion helper
Usage:
  python3 convert.py word    input.pdf output.docx
  python3 convert.py excel   input.pdf output.xlsx
  python3 convert.py encrypt input.pdf output.pdf user_pwd [owner_pwd]
"""
import sys
import os

def to_word(input_pdf, output_docx):
    import fitz  # pymupdf
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc_in = fitz.open(input_pdf)
    doc_out = Document()

    # Set narrow margins
    from docx.shared import Inches
    for section in doc_out.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    for page_num, page in enumerate(doc_in):
        if page_num > 0:
            doc_out.add_page_break()

        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:  # skip image blocks
                continue
            for line in block["lines"]:
                line_text = ""
                is_bold = False
                is_large = False
                for span in line["spans"]:
                    line_text += span["text"]
                    if "bold" in span["font"].lower() or span["flags"] & 2**4:
                        is_bold = True
                    if span["size"] > 13:
                        is_large = True

                line_text = line_text.strip()
                if not line_text:
                    continue

                para = doc_out.add_paragraph()
                run = para.add_run(line_text)
                run.bold = is_bold
                run.font.size = Pt(12 if is_large else 10)

    doc_in.close()
    doc_out.save(output_docx)
    print(f"OK: {output_docx}")

def to_excel(input_pdf, output_xlsx):
    # Use pdfplumber to extract tables, write to xlsx with openpyxl
    import pdfplumber
    import openpyxl
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    with pdfplumber.open(input_pdf) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            ws = wb.create_sheet(title=f"Page {i+1}")
            if tables:
                for table in tables:
                    for row in table:
                        ws.append([cell or '' for cell in row])
                    ws.append([])  # blank row between tables
            else:
                # No tables — extract raw text into column A
                text = page.extract_text() or ''
                for line in text.split('\n'):
                    ws.append([line])

    if not wb.sheetnames:
        wb.create_sheet("Sheet1")

    wb.save(output_xlsx)
    print(f"OK: {output_xlsx}")

def encrypt_pdf(input_pdf, output_pdf, user_pwd, owner_pwd):
    import pikepdf
    with pikepdf.open(input_pdf) as pdf:
        pdf.save(output_pdf, encryption=pikepdf.Encryption(
            user=user_pwd, owner=owner_pwd, R=6
        ))
    print(f"OK: {output_pdf}")

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: convert.py <word|excel|encrypt> <input.pdf> <out> [user_pwd [owner_pwd]]", file=sys.stderr)
        sys.exit(1)

    mode, inp, out = sys.argv[1], sys.argv[2], sys.argv[3]

    if not os.path.exists(inp):
        print(f"Input file not found: {inp}", file=sys.stderr)
        sys.exit(1)

    try:
        if mode == 'word':
            to_word(inp, out)
        elif mode == 'excel':
            to_excel(inp, out)
        elif mode == 'encrypt':
            u = sys.argv[4] if len(sys.argv) > 4 else ''
            o = sys.argv[5] if len(sys.argv) > 5 else u
            encrypt_pdf(inp, out, u, o)
        else:
            print(f"Unknown mode: {mode}", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
